# document_service.py
"""
Document generation service — produces DOCX and PDF.
Single source of truth: SQLite DB (via history items).
notes.txt is no longer used for PDF/DOCX generation.
"""
import logging
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Preformatted, HRFlowable,
)

from utils.file_handler import NOTES_DOCX, NOTES_PDF

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Line classification helpers
# ---------------------------------------------------------------------------

def _is_section_header(line: str) -> bool:
    s = line.strip()
    return s.startswith("===") and s.endswith("===")

def _is_code_fence(line: str) -> bool:
    return line.strip().startswith("```")

def _is_bullet(line: str) -> bool:
    s = line.strip()
    return s.startswith("- ") or s.startswith("* ")

def _is_numbered(line: str) -> bool:
    s = line.strip()
    return len(s) > 2 and s[0].isdigit() and s[1] in ".)"

def _is_arrow_line(line: str) -> bool:
    return line.strip().startswith("→")


# ---------------------------------------------------------------------------
# Append (kept for logging purposes only — DB is source of truth)
# ---------------------------------------------------------------------------

def append_notes(problem_number: str, notes_text: str):
    """No-op kept for import compatibility. DB is now the source of truth."""
    logger.info("append_notes called for %s (DB is source of truth now)", problem_number)


# ---------------------------------------------------------------------------
# Color constants
# ---------------------------------------------------------------------------

PURPLE  = colors.HexColor("#6C63FF")
DARK    = colors.HexColor("#1A1A1A")
LIGHT   = colors.HexColor("#AAAAAA")
CODE_BG = colors.HexColor("#F3F3F3")


# ---------------------------------------------------------------------------
# Safe XML escaping
# ---------------------------------------------------------------------------

def _safe(text: str) -> str:
    return (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )


# ---------------------------------------------------------------------------
# Shared PDF styles
# ---------------------------------------------------------------------------

def _get_styles():
    """Return a dict of all named ParagraphStyles used in PDF generation."""
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "NTitle", parent=base["Title"],
            fontSize=24, textColor=PURPLE, spaceAfter=4, leading=28,
        ),
        "meta": ParagraphStyle(
            "NMeta", parent=base["Normal"],
            fontSize=9, textColor=LIGHT, spaceAfter=16, leading=12,
        ),
        "problem": ParagraphStyle(
            "NProblem", parent=base["Heading2"],
            fontSize=14, textColor=PURPLE,
            spaceBefore=16, spaceAfter=4, leading=18,
            fontName="Helvetica-Bold",
        ),
        "timestamp": ParagraphStyle(
            "NTimestamp", parent=base["Normal"],
            fontSize=8, textColor=LIGHT, spaceAfter=6,
            fontName="Helvetica-Oblique",
        ),
        "section": ParagraphStyle(
            "NSection", parent=base["Heading3"],
            fontSize=11, textColor=DARK,
            spaceBefore=12, spaceAfter=4, leading=14,
            fontName="Helvetica-Bold",
        ),
        "body": ParagraphStyle(
            "NBody", parent=base["Normal"],
            fontSize=10, textColor=DARK, leading=15, spaceAfter=4,
        ),
        "bullet": ParagraphStyle(
            "NBullet", parent=base["Normal"],
            fontSize=10, textColor=DARK, leading=14, spaceAfter=3,
            leftIndent=14,
        ),
        "numbered": ParagraphStyle(
            "NNumbered", parent=base["Normal"],
            fontSize=10, textColor=DARK, leading=14, spaceAfter=3,
            leftIndent=14,
        ),
        "arrow": ParagraphStyle(
            "NArrow", parent=base["Normal"],
            fontSize=9.5, textColor=colors.HexColor("#555555"),
            leading=13, spaceAfter=3, leftIndent=20,
        ),
        "code": ParagraphStyle(
            "NCode", parent=base["Code"],
            fontSize=8.5, fontName="Courier",
            textColor=colors.HexColor("#2D2D2D"),
            backColor=CODE_BG, leading=12,
            leftIndent=10, rightIndent=10,
            spaceBefore=6, spaceAfter=6,
            borderColor=colors.HexColor("#DDDDDD"),
            borderWidth=0.5, borderPadding=6,
        ),
    }


def _build_pdf_cover(st: dict) -> list:
    return [
        Paragraph("DSA Revision Notes", st["title"]),
        Paragraph("AI-Powered Personal Notes · Campus Placement Prep", st["meta"]),
        HRFlowable(width="100%", thickness=1,
                   color=colors.HexColor("#E0E0E0"), spaceAfter=12),
    ]


def _build_problem_header(problem_number: str, timestamp_str: str, st: dict) -> list:
    items = [
        HRFlowable(width="100%", thickness=0.5,
                   color=colors.HexColor("#CCCCCC"), spaceBefore=8),
        Paragraph(f"Problem: {_safe(problem_number)}", st["problem"]),
    ]
    if timestamp_str:
        items.append(Paragraph(f"Generated: {_safe(timestamp_str)}", st["timestamp"]))
    return items


def _line_to_flowables(line: str, st: dict) -> list:
    """Convert a single plain-text line into ReportLab flowables."""
    stripped = line.strip()

    if _is_section_header(line):
        title = stripped.strip("= ").strip()
        return [Spacer(1, 4), Paragraph(f"▸ {_safe(title)}", st["section"])]
    elif _is_arrow_line(stripped):
        return [Paragraph(_safe(stripped), st["arrow"])]
    elif _is_bullet(stripped):
        text = stripped.lstrip("-* ").strip()
        return [Paragraph(f"• {_safe(text)}", st["bullet"])]
    elif _is_numbered(stripped):
        return [Paragraph(_safe(stripped), st["numbered"])]
    elif not stripped:
        return [Spacer(1, 4)]
    else:
        return [Paragraph(_safe(stripped), st["body"])]


def _parse_notes_to_story(notes_text: str, st: dict) -> list:
    """Parse raw notes plain-text into a list of ReportLab flowables."""
    story = []
    lines = notes_text.split("\n")
    in_code = False
    code_buffer: list[str] = []

    def flush_code():
        if not code_buffer:
            return
        story.append(Preformatted("\n".join(code_buffer), st["code"]))
        code_buffer.clear()

    for line in lines:
        if _is_code_fence(line):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buffer.append(line)
            continue
        story += _line_to_flowables(line, st)

    flush_code()
    return story


# ---------------------------------------------------------------------------
# PDF from DB history (primary method)
# ---------------------------------------------------------------------------

def generate_pdf_from_history(history_items: list) -> str:
    """
    Build PDF from DB history items.
    Deleted items are already excluded — caller passes get_history() result.
    Returns path to generated PDF.
    """
    st = _get_styles()

    doc = SimpleDocTemplate(
        NOTES_PDF, pagesize=A4,
        leftMargin=22 * mm, rightMargin=22 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
    )

    story = _build_pdf_cover(st)

    if not history_items:
        story.append(Paragraph("No notes generated yet.", st["body"]))
        doc.build(story)
        return NOTES_PDF

    for item in history_items:
        problem_number = item.get("problem_number", "Unknown")
        notes_text     = item.get("generated_notes", "")
        timestamp      = item.get("timestamp")

        if timestamp:
            try:
                ts_str = datetime.fromtimestamp(float(timestamp)).strftime("%Y-%m-%d %H:%M")
            except Exception:
                ts_str = str(timestamp)
        else:
            ts_str = ""

        story += _build_problem_header(problem_number, ts_str, st)
        story += _parse_notes_to_story(notes_text, st)
        story.append(Spacer(1, 10))

    doc.build(story)
    logger.info("PDF rebuilt from DB history → %s", NOTES_PDF)
    return NOTES_PDF


# ---------------------------------------------------------------------------
# DOCX from DB history (primary method)
# ---------------------------------------------------------------------------

def generate_docx_from_history(history_items: list) -> str:
    """Build DOCX from DB history items. Returns path to generated DOCX."""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Cover
    title = doc.add_heading("DSA Revision Notes", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("AI-Powered Personal Notes · Campus Placement Prep")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(150, 150, 150)
    r.italic = True
    doc.add_paragraph()

    if not history_items:
        doc.add_paragraph("No notes generated yet.")
        doc.save(NOTES_DOCX)
        return NOTES_DOCX

    for item in history_items:
        problem_number = item.get("problem_number", "Unknown")
        notes_text     = item.get("generated_notes", "")
        timestamp      = item.get("timestamp")

        if timestamp:
            try:
                ts_str = datetime.fromtimestamp(float(timestamp)).strftime("%Y-%m-%d %H:%M")
            except Exception:
                ts_str = str(timestamp)
        else:
            ts_str = ""

        # Problem header
        _add_horizontal_rule(doc)
        h = doc.add_heading(f"Problem: {problem_number}", level=2)
        h.runs[0].font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
        if ts_str:
            p = doc.add_paragraph()
            r = p.add_run(f"Generated: {ts_str}")
            r.font.size = Pt(8)
            r.font.color.rgb = RGBColor(160, 160, 160)
            r.italic = True

        # Notes content
        _render_notes_to_docx(doc, notes_text)
        doc.add_paragraph()

    doc.save(NOTES_DOCX)
    logger.info("DOCX rebuilt from DB history → %s", NOTES_DOCX)
    return NOTES_DOCX


def _render_notes_to_docx(doc: Document, notes_text: str):
    """Parse and render raw notes text into the DOCX document."""
    lines = notes_text.split("\n")
    in_code = False
    code_lines: list[str] = []

    def flush_code():
        if not code_lines:
            return
        para = doc.add_paragraph()
        para.paragraph_format.left_indent  = Inches(0.3)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after  = Pt(4)
        run = para.add_run("\n".join(code_lines))
        run.font.name  = "Courier New"
        run.font.size  = Pt(8.5)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F3F3F3")
        pPr.append(shd)
        code_lines.clear()

    for line in lines:
        if _is_code_fence(line):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        stripped = line.strip()

        if _is_section_header(line):
            doc.add_paragraph()
            title = stripped.strip("= ").strip()
            h = doc.add_heading(title, level=3)
            h.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        elif _is_arrow_line(stripped):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(stripped)
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(80, 80, 80)
        elif _is_bullet(stripped):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            p.add_run(stripped.lstrip("-* ").strip()).font.size = Pt(10)
        elif _is_numbered(stripped):
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            p.add_run(stripped[2:].strip()).font.size = Pt(10)
        elif not stripped:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.add_run(stripped).font.size = Pt(10)

    flush_code()


# ---------------------------------------------------------------------------
# Legacy stubs — kept so old imports don't break
# ---------------------------------------------------------------------------

def generate_pdf() -> str:
    """Legacy stub — redirects to generate_pdf_from_history with empty list."""
    from services.db_service import get_history
    items = list(reversed(get_history()))
    return generate_pdf_from_history(items)


def generate_docx() -> str:
    """Legacy stub — redirects to generate_docx_from_history."""
    from services.db_service import get_history
    items = list(reversed(get_history()))
    return generate_docx_from_history(items)


# ---------------------------------------------------------------------------
# DOCX helper
# ---------------------------------------------------------------------------

def _add_horizontal_rule(doc: Document):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para